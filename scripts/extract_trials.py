# -*- coding: utf-8 -*-
"""批量抽取【招募文件2】下所有 .docx / .doc 文本，输出 JSON 便于人工/AI 再整理"""
import json
import os
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    from docx import Document
except ImportError:
    Document = None


SRC = Path(r"e:/Projects/03_患者招募/招募文件2")
OUT = Path(r"e:/Projects/03_患者招募/scripts/招募文件2_raw.json")

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_doc_as_zip(path: Path) -> str:
    """老版 .doc 有时是 zip+xml 伪装；否则尝试 binary 抽文字"""
    try:
        with zipfile.ZipFile(str(path)) as z:
            for name in z.namelist():
                if name.endswith("document.xml"):
                    raw = z.read(name)
                    root = ET.fromstring(raw)
                    parts = []
                    for t in root.iter(W_NS + "t"):
                        if t.text:
                            parts.append(t.text)
                    return "".join(parts)
    except zipfile.BadZipFile:
        pass
    # 二进制 .doc：粗暴抽中文段
    data = path.read_bytes()
    text = ""
    try:
        text = data.decode("utf-16-le", errors="ignore")
    except Exception:
        pass
    # 只保留可打印的汉字+常见符号
    cleaned = []
    for ch in text:
        code = ord(ch)
        if 0x4E00 <= code <= 0x9FFF:
            cleaned.append(ch)
        elif ch in "，。、；：！？·（）《》【】〔〕［］｛｝—·.,;:!?()[]{}<>“”‘’\"' \n\t0123456789+-/\\":
            cleaned.append(ch)
        elif 0x30 <= code <= 0x7E:
            cleaned.append(ch)
        elif ch == "\n":
            cleaned.append(ch)
    out = "".join(cleaned)
    # 压缩长连续空白
    import re
    out = re.sub(r"[ \t]{2,}", " ", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out


def main():
    if Document is None:
        print("python-docx not installed", file=sys.stderr)
        sys.exit(1)
    result = []
    for f in sorted(SRC.iterdir()):
        if not f.is_file():
            continue
        name = f.name
        suffix = f.suffix.lower()
        try:
            if suffix == ".docx":
                text = read_docx(f)
            elif suffix == ".doc":
                text = read_doc_as_zip(f)
            else:
                continue
        except Exception as e:
            text = f"[READ_ERROR] {e!r}"
        result.append({"file": name, "text": text})
        print(f"OK  {name}  {len(text)} chars")
    OUT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n=> {OUT}  ({len(result)} files)")


if __name__ == "__main__":
    main()

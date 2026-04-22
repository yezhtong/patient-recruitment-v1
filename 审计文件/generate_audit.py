"""生成九泰临研 V1 项目立项审计文件（Word 格式）"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 全局页面设置 ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 样式辅助函数 ──────────────────────────────────────────────────────────────
def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=(31, 73, 125))
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=(31, 73, 125))
    elif level == 3:
        set_run_font(run, size=12, bold=True, color=(54, 95, 145))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(20)
    if indent:
        p.paragraph_format.left_indent = Cm(0.74)
    return p

def add_field_row(doc, label, value):
    """单行 label: value 段落"""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}：")
    set_run_font(r1, size=11, bold=True)
    r2 = p.add_run(value if value else "（待填写）")
    set_run_font(r2, size=11)
    if not value:
        r2.font.color.rgb = RGBColor(150, 150, 150)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(20)

def shade_table_header(row, color="1F497D"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    # header
    hdr = table.rows[0]
    shade_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val) if val else "—")
            set_run_font(run, size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # col widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
#  封面
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph("\n\n")
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run("软件研发项目立项审计报告")
r.font.name = "黑体"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = RGBColor(31, 73, 125)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("九泰临研 V1 · 临床试验患者招募平台")
set_run_font(r2, size=16, bold=True, color=(54, 95, 145))

doc.add_paragraph()
proj_info = [
    ("项目名称", "九泰临研患者招募平台 V1"),
    ("所属单位", "九泰药械（医疗器械 CRO）"),
    ("项目负责人", ""),
    ("报告编制人", ""),
    ("报告版本", "V1.0"),
    ("编制日期", "2026 年 4 月 22 日"),
    ("项目状态", "开发中（M0–M6 已完成，M7 部署迁移待启动）"),
]
for label, val in proj_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}：{val if val else '_______________'}")
    set_run_font(r, size=12)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  目录（手动）
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "目  录", level=1)
toc_items = [
    ("一、", "项目概述"),
    ("二、", "立项背景与必要性"),
    ("三、", "项目目标与范围"),
    ("四、", "技术方案"),
    ("五、", "开发进度与里程碑"),
    ("六、", "功能模块清单"),
    ("七、", "人力与资源投入"),
    ("八、", "项目经费预算"),
    ("九、", "风险评估与控制"),
    ("十、", "合规与知识产权"),
    ("十一、", "验收标准"),
    ("十二、", "后续规划"),
    ("附件A", "核心数据模型说明"),
    ("附件B", "路由与权限矩阵"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(f"{num}  {title}")
    set_run_font(r, size=11)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  一、项目概述
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "一、项目概述", level=1)
add_hr(doc)
fields_overview = [
    ("项目全称", "九泰临研患者招募平台 V1"),
    ("英文名称", "JiuTai Clinical Research Patient Recruitment Platform V1"),
    ("项目编号", ""),
    ("立项日期", "2026 年 4 月 18 日"),
    ("预计上线日期", "2026 年（M7 部署迁移完成后，具体日期待定）"),
    ("项目类型", "自主研发软件平台"),
    ("所属业务线", "临床试验患者招募（CRO 增值服务）"),
    ("主要用户群", "潜在受试患者（C 端）、内部运营人员（B 端）"),
    ("项目负责人", ""),
    ("技术负责人", ""),
    ("联系方式", ""),
]
for label, val in fields_overview:
    add_field_row(doc, label, val)

add_body(doc, "\n项目简述：", )
add_body(doc,
    "九泰临研 V1 是面向医疗器械临床试验的一站式患者招募网站，包含患者公开网站（9 个路由）"
    "与统一运营后台（29 个路由）。平台具备试验公示、在线预筛、短信验证注册、报名管理、"
    "病友社区、FAQ 信息中心以及 KPI 运营看板等核心功能，同时满足 Schema.org 医疗结构化数据、"
    "SEO 优化与合规审计要求。"
)

# ══════════════════════════════════════════════════════════════════════════════
#  二、立项背景与必要性
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "二、立项背景与必要性", level=1)
add_hr(doc)
add_heading(doc, "2.1 行业背景", level=2)
add_body(doc,
    "临床试验患者招募是 CRO 业务中耗时最长、成本最高的环节之一。传统招募依赖线下张贴广告、"
    "电话外呼及医生转介等方式，周期长、覆盖面窄、数据留存困难，严重制约项目交付效率。"
)
add_body(doc,
    "据行业数据，约 80% 的临床试验因患者招募不足而延期，平均单个试验招募成本超过 5,000 元/人。"
    "数字化招募平台已成为头部 CRO 企业核心竞争力之一。"
)

add_heading(doc, "2.2 企业内部需求", level=2)
items = [
    "现有招募流程全部依赖线下与电话，无系统化线索管理，易丢单；",
    "试验信息分散于 PDF 与群聊，患者获取渠道不统一，质量参差不齐；",
    "运营缺乏数据看板，无法实时掌握各试验招募漏斗情况；",
    "监管要求临床广告留存电子化记录（广告版本号、伦理批件编号），人工管理合规风险高；",
    "竞品已完成数字化布局，企业面临客户流失压力。",
]
for item in items:
    add_body(doc, f"• {item}", indent=True)

add_heading(doc, "2.3 立项必要性结论", level=2)
add_body(doc,
    "综上，自主研发患者招募平台能有效降低招募成本、缩短招募周期、提升合规管理水平，"
    "具备明确的业务价值与商业可行性，建议予以立项。"
)

# ══════════════════════════════════════════════════════════════════════════════
#  三、项目目标与范围
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "三、项目目标与范围", level=1)
add_hr(doc)
add_heading(doc, "3.1 核心目标", level=2)
goals = [
    ("G1", "上线患者公开网站，支持试验信息公示与在线预筛，预计将单项目招募响应率提升 30% 以上"),
    ("G2", "打通「预筛提交 → 线索入池 → 运营跟进 → 报名管理」全链路，消除信息孤岛"),
    ("G3", "建立病友社区，提升患者粘性与自发传播，降低获客成本"),
    ("G4", "构建运营看板与审计日志，满足监管合规与内部管理需求"),
    ("G5", "技术架构支持未来扩展（多试验、多站点、多语言）"),
]
add_table(doc, ["目标编号", "描述"], goals, [2, 13])

add_heading(doc, "3.2 V1 版本范围", level=2)
scope_in = [
    "患者端：首页、试验列表、试验详情、在线预筛、成功页、登录/注册（手机验证码）、我的报名、常见问题、病友社区（6 个分区）",
    "运营后台：试验管理 CRUD、线索池、报名管理、社区审核、KPI 看板、CSV 导出、FAQ/分区/敏感词配置、审计日志",
    "SEO 与合规：Schema.org JSON-LD、sitemap、robots、广告版本与伦理批件字段",
    "技术底座：Next.js 16 App Router、Prisma 7 ORM、SQLite（开发）/ MySQL-Postgres（生产）、iron-session 鉴权",
]
add_heading(doc, "范围内（In-Scope）：", level=3)
for item in scope_in:
    add_body(doc, f"✓ {item}", indent=True)

scope_out = [
    "支付功能",
    "电子病历 / HIS 接口对接",
    "多语言国际化",
    "移动端 App",
    "第三方 BI 集成",
]
add_heading(doc, "范围外（Out-of-Scope）：", level=3)
for item in scope_out:
    add_body(doc, f"✗ {item}", indent=True)

# ══════════════════════════════════════════════════════════════════════════════
#  四、技术方案
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "四、技术方案", level=1)
add_hr(doc)
add_heading(doc, "4.1 技术栈选型", level=2)
tech = [
    ("前端框架", "Next.js 16.2（App Router + React 19 + Turbopack）", "全栈一体、SSR/SSG 支持 SEO、TypeScript 强类型"),
    ("语言", "TypeScript 5", "编译期错误捕获，降低运行时 bug"),
    ("样式", "Tailwind CSS v4 + 模块化 CSS", "快速迭代，设计 Token 统一管理"),
    ("ORM", "Prisma 7", "类型安全查询，一行切换 SQLite / MySQL / Postgres"),
    ("开发数据库", "SQLite via @prisma/adapter-better-sqlite3", "零配置本机开发，无需额外服务"),
    ("生产数据库", "MySQL 或 PostgreSQL（M7 切换）", "高并发稳定性，支持数据备份"),
    ("鉴权", "iron-session（HttpOnly Cookie）", "轻量无状态，管理员/患者双 session 隔离"),
    ("数据校验", "Zod", "运行时 schema 校验，与 TypeScript 类型自动同步"),
    ("密码", "Node 原生 scrypt", "无第三方依赖，符合 NIST 推荐"),
    ("部署", "Docker + docker-compose（M7）", "一键部署，环境一致性"),
    ("短信", "SmsProvider 接口 + DevConsoleSmsProvider（开发）", "可插拔设计，生产只需替换实现类"),
]
add_table(doc, ["技术组件", "选型", "选型理由"], tech, [3, 5.5, 6.5])

add_heading(doc, "4.2 架构说明", level=2)
add_body(doc,
    "采用前后端一体的 Next.js App Router 架构：患者端页面通过 Server Component 直接读取数据库"
    "（支持 SSR/ISR SEO），交互表单通过 Server Action 校验写入，无需独立 REST API 层。"
    "运营后台同架构，采用独立的 iron-session cookie（`jt_admin_session`）与路由守卫隔离。"
    "短信、敏感词、CSV 工具均封装为可插拔 lib 模块，具备良好的可替换性。"
)

# ══════════════════════════════════════════════════════════════════════════════
#  五、开发进度与里程碑
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "五、开发进度与里程碑", level=1)
add_hr(doc)
milestones = [
    ("M0", "环境与骨架", "2026-04-18", "2026-04-18", "✅ 完成", "Next.js + Prisma + SQLite 三条腿跑通，Git 初始化"),
    ("M1", "9 个患者端页面迁移", "2026-04-19", "2026-04-19", "✅ 完成", "9 路由全部 200，TS 零错误，样式贴近原型"),
    ("M1.5", "样式校准", "2026-04-19", "2026-04-19", "✅ 完成", "9 个 page-scoped CSS 搬入，视觉对齐原型"),
    ("M1.6", "真实招募数据接入", "2026-04-19", "2026-04-19", "✅ 完成", "解析 3 份真实招募广告，种子数据写入 DB"),
    ("M2", "主链路闭环 + 运营后台 CRUD", "2026-04-19", "2026-04-19", "✅ 完成", "预筛真入库、Lead 池、试验 CRUD、admin 登录守卫，15 路由"),
    ("M3", "账号真接入 + 我的报名", "2026-04-19", "2026-04-19", "✅ 完成", "手机验证码登录、User/SmsCode/Application 模型，17 路由"),
    ("M4", "信任层 + SEO", "2026-04-19", "2026-04-19", "✅ 完成", "FAQ 20 条、JSON-LD、sitemap、robots，19 路由"),
    ("M5", "病友社区模块", "2026-04-19", "2026-04-19", "✅ 完成", "6 分区 + 发帖/评论 + 敏感词 + 匿名 + 审核后台，25 路由"),
    ("M6", "运营完善", "2026-04-20", "2026-04-20", "✅ 完成", "权限审计 + KPI 看板 + CSV 导出 + FAQ/分区/敏感词后台，38 路由"),
    ("M7", "部署迁移", "待定", "待定", "⏳ 待启动", "Docker 化、SQLite→MySQL/Postgres、SMS 真接、SESSION 密钥"),
]
add_table(doc,
    ["编号", "名称", "开始", "完成", "状态", "主要产出"],
    milestones,
    [1, 3.5, 2, 2, 2, 5.5]
)
add_body(doc,
    "注：M0–M6 累计开发周期约 3 天（2026-04-18 至 2026-04-20），采用 AI 辅助编程方式高效交付。"
    "M7 部署迁移待 IT 确认服务器配置后启动。"
)

# ══════════════════════════════════════════════════════════════════════════════
#  六、功能模块清单
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "六、功能模块清单", level=1)
add_hr(doc)
add_heading(doc, "6.1 患者公开网站（9 个路由）", level=2)
patient_routes = [
    ("/", "网站首页", "试验分类导览、机构介绍、快捷入口"),
    ("/trials", "试验列表", "按疾病类型/状态筛选，分页展示"),
    ("/trials/[slug]", "试验详情", "完整试验信息 + Schema.org MedicalStudy JSON-LD + 一键拨号"),
    ("/prescreen/[slug]", "在线预筛", "定制化问卷表单，提交写入 Lead 并创建 Application"),
    ("/prescreen/success", "预筛成功", "生成唯一提交编号 PS-YYYYMMDD-XXXXXX"),
    ("/auth", "登录/注册", "手机验证码 OTP，60 秒节流，30 天 session"),
    ("/me", "我的报名", "展示当前账号关联的所有 Application，支持阶段查看"),
    ("/faq", "常见问题", "20 条 FAQ 分类展示，FAQPage JSON-LD"),
    ("/community", "病友社区", "6 个病种分区，帖子/评论，匿名发帖，敏感词过滤"),
]
add_table(doc, ["路由", "页面名称", "核心功能"], patient_routes, [3.5, 3, 9.5])

add_heading(doc, "6.2 运营后台（29 个路由）", level=2)
admin_routes = [
    ("/admin", "工作台首页", "KPI 看板：6 指标卡 + SVG 漏斗 + 堆叠条 + Top5 + 待办"),
    ("/admin/trials", "试验管理列表", "搜索/筛选，线索计数"),
    ("/admin/trials/new", "新建试验", "27 字段完整表单，含伦理批件与广告版本号"),
    ("/admin/trials/[id]/edit", "编辑/删除试验", "仅 admin 角色可操作，写审计日志"),
    ("/admin/leads", "线索池", "姓名/手机号/状态/试验筛选，手机号脱敏展示"),
    ("/admin/leads/[id]", "线索详情", "基础资料 + 预筛答案 + 合规授权 + 状态流转 + 备注"),
    ("/admin/applications", "报名管理", "阶段/试验筛选，手机号脱敏"),
    ("/admin/applications/[id]", "报名详情", "患者账户 + 试验 + 预筛答案 + 阶段流转 + 备注"),
    ("/admin/community/posts", "社区帖子审核", "状态/分区/风险筛选，命中计数"),
    ("/admin/community/posts/[id]", "帖子审核详情", "全文 + 审核日志 + 5 个操作（approve/reject/hide/feature）"),
    ("/admin/community/sensitive-hits", "敏感词命中统计", "按 keyword × riskType × riskLevel groupBy"),
    ("/admin/community/groups", "社区分区配置", "启停 toggle + 帖子数 + 新建/编辑/删除（有帖子保护）"),
    ("/admin/community/sensitive-words", "敏感词库管理", "增删改 + 批量导入（level|type|keyword 格式）"),
    ("/admin/faq", "FAQ 管理", "分类/状态/搜索 + 排序，一键 toggle 发布"),
    ("/admin/audit-logs", "审计日志", "最近 200 条，支持按 entityType/action/关键词筛选"),
    ("/api/admin/export/leads", "线索 CSV 导出", "16 字段含手机号明文，按筛选条件导出，仅 admin"),
    ("/api/admin/export/applications", "报名 CSV 导出", "17 字段含手机号明文，仅 admin，写审计"),
]
add_table(doc, ["路由", "页面/接口名称", "核心功能"], admin_routes, [4, 3.5, 8.5])

add_heading(doc, "6.3 SEO 与合规基础设施", level=2)
seo_items = [
    ("/sitemap.xml", "动态站点地图，含静态路由 + 所有公开试验 + 社区分区与帖子（上限 500）"),
    ("/robots.txt", "爬虫规则，禁止索引 /admin/ /me /prescreen/，社区公开可索引"),
    ("JSON-LD", "试验详情页 MedicalStudy、首页 MedicalOrganization、FAQ 页 FAQPage、社区帖子 DiscussionForumPosting"),
    ("广告合规字段", "ClinicalTrial 模型内置 adVersion / adVersionDate / ethicsApproval，每次变更写审计日志"),
]
for route, desc in seo_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{route}：")
    set_run_font(r1, size=11, bold=True)
    r2 = p.add_run(desc)
    set_run_font(r2, size=11)
    p.paragraph_format.space_after = Pt(3)

# ══════════════════════════════════════════════════════════════════════════════
#  七、人力与资源投入
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "七、人力与资源投入", level=1)
add_hr(doc)
add_heading(doc, "7.1 开发团队", level=2)
team = [
    ("产品经理", "1", "需求设计、PRD 编写、验收把关"),
    ("全栈工程师（AI 辅助）", "1", "Next.js 前后端、Prisma 数据层、Server Action、鉴权、SEO"),
    ("UI/UX 设计师", "1", "设计系统维护、页面视觉规范、组件规格"),
    ("质量工程师", "1", "TypeScript 编译验收、构建验收、路由抽检、E2E 脚本"),
    ("运维工程师", "1", "Docker 化、生产数据库迁移、环境变量与密钥管理（M7）"),
]
add_table(doc, ["角色", "人数", "主要职责"], team, [4, 1.5, 10.5])

add_heading(doc, "7.2 开发工时估算", level=2)
effort = [
    ("M0–M1.6", "环境搭建 + 页面迁移 + 真实数据接入", "约 8 人时"),
    ("M2", "主链路闭环 + 运营后台 CRUD", "约 10 人时"),
    ("M3", "账号体系 + 短信 + Application 模型", "约 8 人时"),
    ("M4", "FAQ + SEO + JSON-LD + sitemap", "约 6 人时"),
    ("M5", "病友社区（发帖/评论/审核/敏感词）", "约 14 人时"),
    ("M6", "审计底座 + KPI 看板 + 导出 + 4 个后台配置模块", "约 16 人时"),
    ("M7", "Docker + DB 迁移 + SMS 真接（待开始）", "预估约 12 人时"),
    ("合计（M0–M6）", "", "约 62 人时"),
]
add_table(doc, ["里程碑", "工作内容", "工时估算"], effort, [2.5, 8, 5.5])
add_body(doc, "注：以上工时为估算值，实际投入以打卡记录与 Git 提交时间戳为准。")

add_heading(doc, "7.3 硬件资源", level=2)
hw = [
    ("开发机", "Windows 11 Pro，本地运行 Next.js dev server + SQLite"),
    ("生产服务器", "公司 Linux 服务器（发行版待 IT 确认），M7 部署时使用"),
    ("域名 / SSL", "待申请，上线前需配置 HTTPS"),
    ("短信通道", "待选型（2026-04-25 前询价），开发期使用固定验证码 123456"),
]
add_table(doc, ["资源类型", "说明"], hw, [3, 13])

# ══════════════════════════════════════════════════════════════════════════════
#  八、项目经费预算
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "八、项目经费预算", level=1)
add_hr(doc)
add_body(doc, "注：以下为参考预算框架，具体金额请财务与采购部门根据实际情况填写。")
budget = [
    ("人力成本", "研发人员工资 / 外包费用", "", ""),
    ("服务器成本", "Linux 服务器（含 IDC 或云）年费", "", ""),
    ("域名 / SSL 证书", "年费", "", ""),
    ("短信服务费", "按条计费（预估年发送量 × 单价）", "", ""),
    ("软件许可费", "暂无（全部开源技术栈）", "0", ""),
    ("AI 辅助工具费", "Claude Code 等 AI 编程工具订阅", "", ""),
    ("测试与验收", "QA 工时费用", "", ""),
    ("运营维护（首年）", "服务器运维 + bug 修复 + 需求迭代", "", ""),
    ("合计", "", "", ""),
]
add_table(doc,
    ["费用类别", "说明", "预算金额（元）", "备注"],
    budget,
    [3.5, 5.5, 3, 4]
)

# ══════════════════════════════════════════════════════════════════════════════
#  九、风险评估与控制
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "九、风险评估与控制", level=1)
add_hr(doc)
risks = [
    ("R1", "数据丢失", "高", "代码与数据库仅本机一份，无远程备份",
     "立即将 app/ 目录与 dev.db 备份至移动硬盘；M7 上线前配置自动化数据库备份"),
    ("R2", "短信服务未接入", "高", "生产环境验证码仍用固定码 123456，无法真实注册",
     "2026-04-25 前完成服务商询价与接入；SmsProvider 接口已预留，切换成本低"),
    ("R3", "SESSION_PASSWORD 未替换", "高", "dev 硬编码密钥泄露会导致 session 伪造",
     "M7 部署时必须注入真实随机密钥至环境变量，绝不将密钥提交至代码库"),
    ("R4", "敏感词库待审定", "中", "当前词库为 dev 默认 35 条，未经运营审定",
     "上线前由运营团队完成词库审定，后台化管理界面已就绪（/admin/community/sensitive-words）"),
    ("R5", "服务器配置未确认", "中", "M7 依赖 IT 提供服务器发行版、端口等信息",
     "在 M7 启动前 2 周完成服务器信息确认"),
    ("R6", "患者数据隐私", "中", "手机号、预筛答案为敏感个人信息",
     "导出接口仅限 admin 角色；生产环境启用 HTTPS；未来可评估加密存储方案"),
    ("R7", "临床广告合规", "中", "试验广告须符合 GCP 及药监局要求",
     "系统已内置广告版本号与伦理批件字段，每次修改写审计日志，供监管核查"),
]
add_table(doc,
    ["编号", "风险描述", "等级", "风险原因", "控制措施"],
    risks,
    [1, 3, 1.5, 4, 6.5]
)

# ══════════════════════════════════════════════════════════════════════════════
#  十、合规与知识产权
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "十、合规与知识产权", level=1)
add_hr(doc)
add_heading(doc, "10.1 个人信息保护", level=2)
add_body(doc,
    "平台收集患者手机号及预筛健康信息，属于个人敏感信息，须遵循《个人信息保护法》相关规定。"
    "系统已内置隐私协议授权字段（`agreePrivacy`）并在预筛提交时强制勾选。"
    "手机号在运营后台列表页脱敏展示，仅在导出 CSV 时向 admin 角色提供明文，每次导出写审计日志。"
)

add_heading(doc, "10.2 临床试验信息合规", level=2)
add_body(doc,
    "试验信息模型包含广告版本号（`adVersion`）、广告版本日期（`adVersionDate`）"
    "与伦理批件编号（`ethicsApproval`）字段，试验信息每次新建或修改均写入 `AuditLog`，"
    "满足临床试验广告存档与变更追溯要求。"
)

add_heading(doc, "10.3 开源许可证", level=2)
open_source = [
    ("Next.js", "MIT"),
    ("React", "MIT"),
    ("Prisma", "Apache 2.0"),
    ("TypeScript", "Apache 2.0"),
    ("Tailwind CSS", "MIT"),
    ("iron-session", "MIT"),
    ("Zod", "MIT"),
    ("better-sqlite3", "MIT"),
    ("python-docx（报告生成工具）", "MIT"),
]
add_table(doc, ["依赖库", "许可证"], open_source, [6, 10])
add_body(doc, "以上开源许可证均允许商业用途，不构成知识产权风险。")

add_heading(doc, "10.4 知识产权归属", level=2)
add_field_row(doc, "代码版权归属", "")
add_field_row(doc, "是否申请软件著作权", "")
add_field_row(doc, "申请日期", "")

# ══════════════════════════════════════════════════════════════════════════════
#  十一、验收标准
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "十一、验收标准", level=1)
add_hr(doc)
add_heading(doc, "11.1 技术验收（已达成，M0–M6）", level=2)
tech_verify = [
    ("TypeScript 编译", "npx tsc --noEmit 零错误", "✅ 已通过"),
    ("生产构建", "npm run build 成功，38 路由全部编译", "✅ 已通过"),
    ("路由可访问性", "所有公开路由 HTTP 200，admin 路由未登录 307", "✅ 已通过"),
    ("权限守卫", "admin 写操作仅 admin 角色可访问，导出 API 未授权返回 401", "✅ 已通过"),
    ("数据库迁移", "prisma migrate deploy 全部通过", "✅ 已通过"),
]
add_table(doc, ["验收项", "标准", "状态"], tech_verify, [3.5, 9, 3.5])

add_heading(doc, "11.2 功能验收（待 M7 上线后完整验收）", level=2)
func_verify = [
    ("患者端全链路", "预筛提交 → Lead 入库 → 短信验证码 → 登录 → 我的报名真数据", "待真 SMS 接入后验收"),
    ("运营后台管理", "试验 CRUD / 线索状态流转 / 报名阶段推进 / 导出 CSV", "开发环境已验收"),
    ("社区功能", "发帖/评论/匿名/敏感词拦截/admin 审核", "开发环境已验收"),
    ("KPI 看板", "6 指标卡数据准确，SVG 漏斗计算正确", "开发环境已验收"),
    ("SEO", "sitemap.xml / robots.txt 可访问，JSON-LD 格式正确", "开发环境已验收"),
    ("性能", "关键页面首屏时间 < 3 秒（生产环境）", "待生产环境测试"),
    ("安全", "XSS / SQL 注入 / CSRF 防护验证", "待生产环境安全扫描"),
]
add_table(doc, ["验收项", "验收标准", "当前状态"], func_verify, [3.5, 7, 5.5])

# ══════════════════════════════════════════════════════════════════════════════
#  十二、后续规划
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "十二、后续规划", level=1)
add_hr(doc)
add_heading(doc, "12.1 M7 部署迁移（近期）", level=2)
m7_items = [
    "Dockerfile + docker-compose 编写与调试",
    "SQLite → MySQL / PostgreSQL 数据库迁移脚本",
    "SESSION_PASSWORD 切换为真实随机密钥",
    "DevConsoleSmsProvider 替换为真实 SMS 服务商（2026-04-25 前完成询价）",
    "NEXT_PUBLIC_SITE_URL 等环境变量生产化配置",
    "og 图片与客服联系方式正式值替换",
]
for item in m7_items:
    add_body(doc, f"• {item}", indent=True)

add_heading(doc, "12.2 V2 规划方向（中期）", level=2)
v2_items = [
    "多试验多站点支持（SaaS 化改造）",
    "患者端微信小程序",
    "与 HIS 系统接口对接（条件成熟时）",
    "数据分析模块增强（患者画像、招募预测）",
    "多语言支持（英文版）",
]
for item in v2_items:
    add_body(doc, f"• {item}", indent=True)

add_heading(doc, "12.3 待人工确认事项", level=2)
pending = [
    ("短信服务商接入", "2026-04-25 前", "负责人"),
    ("敏感词库正式审定", "M7 部署前", "运营团队"),
    ("服务器发行版与部署凭证确认", "M7 启动前 2 周", "IT 部门"),
    ("og 图片与客服联系方式正式值", "上线前", "运营团队"),
    ("软件著作权申请", "建议尽快启动", "法务/知识产权部门"),
    ("域名申请与 SSL 配置", "M7 前", "IT 部门"),
]
add_table(doc, ["待确认事项", "截止时间", "责任人"], pending, [6, 3.5, 6.5])

# ══════════════════════════════════════════════════════════════════════════════
#  附件 A：核心数据模型说明
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "附件 A：核心数据模型说明", level=1)
add_hr(doc)
models = [
    ("ClinicalTrial", "临床试验",
     "slug, title, phase, status, sponsor, intervention, studyDesign, targetEnrollment, "
     "siteName, siteAddress, contactPerson, contactPhone, benefits, adVersion, "
     "adVersionDate, ethicsApproval, qrcodeUrl（共 27+ 字段）"),
    ("Lead", "预筛线索",
     "trialId, userId（可空）, name, phone, gender, age, city, projectAnswers（JSON）, "
     "agreePrivacy, sourcePage, status（new/contacted/qualified/disqualified）, note"),
    ("User", "患者账号",
     "phone（唯一）, name, gender, age, city, agreeReuse, lastLoginAt"),
    ("SmsCode", "短信验证码",
     "phone, code, expiresAt（10 分钟）, consumedAt, attempts（上限 5 次）"),
    ("Application", "报名记录",
     "userId, trialId, leadId（可空）, stage（submitted/in_review/contacted/enrolled/withdrawn/closed）, nextAction, note"),
    ("OperatorUser", "后台操作员",
     "username, passwordHash（scrypt）, role（admin/operator）, status"),
    ("FaqArticle", "常见问题",
     "slug, question, answer, category, tags, order, isPublished"),
    ("CommunityGroup", "社区分区",
     "slug, name, diseaseTag, introduction, isEnabled, sortOrder"),
    ("CommunityPost", "社区帖子",
     "groupId, authorUserId, isAnonymous, title, content, postType, "
     "reviewStatus（pending/approved/rejected/hidden/featured）, relatedTrialId"),
    ("CommunityComment", "社区评论",
     "postId, authorUserId, isAnonymous, reviewStatus"),
    ("SensitiveWord", "敏感词",
     "keyword, type（contact/scam/medical-claim/privacy/ad）, level（high/medium/low）"),
    ("AuditLog", "操作审计日志",
     "operatorId, action, entityType, entityId, summary, detail, createdAt"),
]
add_table(doc, ["模型名称", "中文名", "主要字段"], models, [3.5, 2.5, 10])

# ══════════════════════════════════════════════════════════════════════════════
#  附件 B：路由与权限矩阵
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "附件 B：路由与权限矩阵", level=1)
add_hr(doc)
perm_matrix = [
    ("患者公开网站", "/  /trials  /faq  /community  /about  /contact", "所有人", "否"),
    ("试验详情", "/trials/[slug]", "所有人", "否"),
    ("在线预筛", "/prescreen/[slug]", "所有人（不登录也可提交）", "否"),
    ("登录/注册", "/auth", "未登录用户", "否"),
    ("我的报名", "/me", "已登录患者（jt_user_session）", "是 → /auth"),
    ("社区发帖", "/community/new", "已登录患者", "是 → /auth"),
    ("后台所有页面", "/admin/*", "已登录 operator/admin（jt_admin_session）", "是 → /admin/login"),
    ("试验 CRUD", "/admin/trials/new  /admin/trials/[id]/edit", "仅 admin 角色", "是 → 403/redirect"),
    ("审计日志", "/admin/audit-logs", "仅 admin 角色", "是 → redirect"),
    ("CSV 导出", "/api/admin/export/*", "仅 admin 角色", "是 → 401"),
]
add_table(doc,
    ["功能区", "路由范围", "可访问角色", "未授权处理"],
    perm_matrix,
    [2.5, 5.5, 4, 4]
)

# ══════════════════════════════════════════════════════════════════════════════
#  页脚
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run("本报告由项目组编制，如有疑问请联系项目负责人。编制日期：2026 年 4 月 22 日")
set_run_font(r, size=9, color=(128, 128, 128))

# ── 保存 ──────────────────────────────────────────────────────────────────────
output_path = r"e:\Projects\03_患者招募\审计文件\九泰临研V1_项目立项审计报告_2026-04-22.docx"
doc.save(output_path)
print(f"已生成：{output_path}")
